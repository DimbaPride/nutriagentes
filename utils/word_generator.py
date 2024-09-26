from docx import Document

def generate_word_doc(student_name, diet_plan, workout_plan):
    doc = Document()
    
    # Título do documento
    doc.add_heading(f'Plano de Dieta e Treino - {student_name}', 0)
    
    # Seção da dieta
    doc.add_heading('Plano de Dieta', level=1)
    doc.add_paragraph(diet_plan)
    
    # Seção do treino
    doc.add_heading('Plano de Treino', level=1)
    doc.add_paragraph(workout_plan)
    
    # Salvando o documento
    file_name = f'{student_name}_Plano_Dieta_Treino.docx'
    doc.save(file_name)
    
    return file_name
